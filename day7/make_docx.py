# B-07-O2 — PRD lite 기반 사업 소개 DOCX 자동 생성 스크립트
# 근거 자료: B-02-O1(PRD lite) · B-02-R1(시장조사) · B-02-R2(경쟁사) · B-07-O1(흐름도)
# 실행: .venv/bin/python day7/make_docx.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 한글 폰트 설정 (본문 스타일) ──
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

BROWN = RGBColor(0x5C, 0x3D, 0x2E)
ACCENT = RGBColor(0xC4, 0x71, 0x3B)

def set_korean(run):
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def heading(text, level=1):
    h = doc.add_heading('', level=level)
    r = h.add_run(text)
    set_korean(r)
    r.font.color.rgb = BROWN if level == 1 else ACCENT
    return h

def para(text, bold=False, size=None, align=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean(r)
    r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 6'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_korean(r); r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(v)
            set_korean(r)
    return t

def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ═══════ ① 표지 ═══════
for _ in range(6): doc.add_paragraph()
para('사업 소개서', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=ACCENT)
para('모닝브루', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, color=BROWN)
para('망원동 핸드드립 카페 · 아침 픽업 예약 서비스', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('대표  정해린', align=WD_ALIGN_PARAGRAPH.CENTER)
para('2026년 7월 14일', align=WD_ALIGN_PARAGRAPH.CENTER)
para('(바이브 코딩 기초반 실습용 가상 사업 문서)', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
page_break()

# ═══════ 목차 (수동 — 모든 뷰어에서 동일 렌더링) ═══════
# ※ Word 전용 TOC 필드는 Pages 등 다른 뷰어에서 줄바꿈이 깨져 수동 목차로 대체
para('목차', bold=True, size=18, color=BROWN)
doc.add_paragraph()
for entry, sub in [
    ('1. 사업 개요', ['1.1 문제 (가설 — 검증 진행 중)', '1.2 타깃', '1.3 가치 — 차별화 가설']),
    ('2. 시장 현황', ['2.1 직접 경쟁 (주변 카페)', '2.2 대체재 (카페 밖)']),
    ('3. 서비스 소개', ['3.1 핵심 흐름', '3.2 운영 방식 (실제 구현 기준)']),
    ('4. 연락처 및 다음 단계', ['4.1 연락처', '4.2 다음 단계']),
]:
    para(entry, bold=True, size=12)
    for sb in sub:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(sb)
        set_korean(r)
        r.font.size = Pt(10.5)
page_break()

# ═══════ ② 사업 개요 ═══════
heading('1. 사업 개요', 1)
heading('1.1 문제 (가설 — 검증 진행 중)', 2)
para('출근길 직장인은 아침에 커피를 사려면 프랜차이즈 줄에 서야 해 출근 시간에 쫓기고, '
     '대기 시간을 예측할 수 없다고 가정한다. 편의점 커피는 빠르지만 품질이 아쉽다. '
     '이 불편의 실재 여부는 시장조사(B-02-R1) 검증 항목으로 관리하고 있다.')
heading('1.2 타깃', 2)
para('망원역 2번 출구 인근으로 출근하는 30대 직장인 — 아침 8~9시, 시간에 쫓기며 '
     '품질 좋은 커피를 원하는 사람. (모든 고객이 아니라 하나의 고객에 집중)')
heading('1.3 가치 — 차별화 가설', 2)
para('이른 아침(08시 오픈) + 직접 내린 핸드드립 품질 + 골목의 편안함 + 사전 픽업 예약. '
     '네 가지를 동시에 제공하는 경쟁 대안은 확인된 범위에서 드물다 (실제 상권 조사로 검증 예정).')
para('포지셔닝: "줄 서는 5분, 예약하면 0분."', bold=True)
page_break()

# ═══════ ③ 시장 현황 ═══════
heading('2. 시장 현황', 1)
para('※ 아래 경쟁 대안 정보는 일반적 카페 유형에 대한 가정이며, 특정 점포 실사 결과가 아니다. '
     '가격 등 변동 정보는 확인이 필요하다.', size=9)
heading('2.1 직접 경쟁 (주변 카페)', 2)
table(['경쟁 대안', '강점', '약점 (모닝브루의 기회)', '검증 상태'], [
    ['대형 프랜차이즈', '앱 주문·픽업, 브랜드 신뢰', '아침 붐빔·대기, 개성 없음', '일반 사실'],
    ['저가 프랜차이즈', '저렴, 빠름', '품질 편차, 정성 부족', '가정'],
    ['개인 감성 카페', '인테리어·분위기', '늦은 오픈(11시~), 픽업 없음', '가정 — 상권 실사 필요'],
    ['무인 카페', '24시간, 대기 없음', '품질 낮음, 응대 없음', '가정'],
])
heading('2.2 대체재 (카페 밖)', 2)
table(['대체재', '강점', '약점 (모닝브루의 기회)', '검증 상태'], [
    ['편의점 원두커피', '초저가, 어디에나', '맛·경험 최저', '일반 사실'],
    ['회사 탕비실', '무료', '믹스 품질, 기분전환 안 됨', '가정'],
    ['배달 앱 커피', '이동 불필요', '배달비·지연, 식은 커피', '일반 사실'],
])
page_break()

# ═══════ ④ 서비스 소개 ═══════
heading('3. 서비스 소개', 1)
heading('3.1 핵심 흐름', 2)
para('페이지 방문 → 메뉴 확인 → 웹 폼으로 픽업 예약(메뉴·수량·날짜·시간) → 매장에서 픽업. '
     '예약은 1분 이내로 끝나며, 접수 상태에서는 고객이 온라인으로 직접 취소할 수 있다.')
doc.add_picture('day7/B-07-O1.png', width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para('그림 1. 사용자 흐름도 — 진입·행동·결과·분기(오류 경로 포함)', size=9,
     align=WD_ALIGN_PARAGRAPH.CENTER)
heading('3.2 운영 방식 (실제 구현 기준)', 2)
table(['항목', '내용'], [
    ['픽업 시간', '매일 08:00~10:00, 10분 단위 슬롯 (월요일 휴무)'],
    ['슬롯 정책', '품질 유지를 위해 슬롯당 3잔 한정, 마감 슬롯은 선택 차단'],
    ['메뉴·가격', '핸드드립 5,500원 / 오늘의 원두 4,500원 / 수제 스콘 3,500원 / 에이드 6,000원'],
    ['예약 관리', '접수 → 확정 → 픽업완료 상태 관리, 확정 후 삭제 잠금'],
    ['개인정보', '이름·연락처만 수집, 목적·보유기간(1개월) 고지 및 동의'],
])
para('공개 서비스 URL: https://vibe-coding-practice-xi.vercel.app/day1/', size=9)
page_break()

# ═══════ ⑤ 연락처·다음 단계 ═══════
heading('4. 연락처 및 다음 단계', 1)
heading('4.1 연락처', 2)
table(['구분', '내용'], [
    ['매장', '서울 망원동 골목 끝 (망원역 2번 출구 도보 6분)'],
    ['전화', '02-123-4567'],
    ['인스타그램', '@morningbrew_official'],
    ['영업시간', '화~일 08:00~20:00 (월요일 정기 휴무)'],
])
heading('4.2 다음 단계', 2)
for i, step in enumerate([
    '아침 픽업 수요 실측 — 매장 방문객 간이 설문 및 베타 운영 (시장조사 검증 항목)',
    '예약 데이터 서버 수집 체계 도입 (현재 브라우저 저장 → DB 연동 예정)',
    '노쇼 대응 정책 수립 (보관 시간·연락 절차)',
    '검색 노출 강화 — 구조화 데이터·매장 사진·FAQ 추가',
], 1):
    para(f'{i}. {step}')

doc.save('day7/B-07-O2.docx')
print('day7/B-07-O2.docx 생성 완료')
